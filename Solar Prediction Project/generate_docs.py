"""
Generate comprehensive project documentation as a Word (.docx) document.
Covers all modules, HTML, CSS, JS with line-by-line explanations.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import datetime

doc = Document()

# ==============================
# STYLES
# ==============================
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Code style
code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_font = code_style.font
code_font.name = 'Consolas'
code_font.size = Pt(9)
code_font.color.rgb = RGBColor(30, 30, 30)
code_style.paragraph_format.space_before = Pt(2)
code_style.paragraph_format.space_after = Pt(2)

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def add_numbered(text):
    doc.add_paragraph(text, style='List Number')

def add_concept(title, explanation):
    p = doc.add_paragraph()
    run = p.add_run(f'💡 Concept — {title}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 100, 180)
    p.add_run(explanation)

def add_code(code_text):
    """Add a code block with monospace formatting."""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph(style='CodeBlock')
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

# ==============================
# COVER PAGE
# ==============================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('SolarVista — Solar Power Prediction Web App', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Complete Project Documentation')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run(f'Date: {datetime.now().strftime("%B %d, %Y")}\n').font.size = Pt(12)
info.add_run('Version: 2.0\n').font.size = Pt(12)
info.add_run('Tech Stack: Flask, Python, Chart.js, XGBoost\n').font.size = Pt(12)

doc.add_page_break()

# ==============================
# TABLE OF CONTENTS
# ==============================
add_heading('Table of Contents', level=1)
toc_items = [
    '1. Project Overview & Architecture',
    '2. File Structure',
    '3. Data Flow Diagram',
    '4. data_collections.py — Data Gathering',
    '5. prepare_dataset.py — Data Preprocessing',
    '6. train_model.py — ML Model Training',
    '7. predict_solar.py — Standalone Predictions',
    '8. weather_api.py — OpenWeatherMap Integration',
    '9. predictor.py — ML Prediction Module',
    '10. metrics.py — Panel & Environmental Calculations',
    '11. app.py — Flask Backend & Routes',
    '12. index.html — Frontend Layout',
    '13. style.css — Styling & Themes',
    '14. app.js — Frontend JavaScript Logic',
    '15. API Limitations & Assumptions',
    '16. Key Formulas & Calculations',
]
for item in toc_items:
    add_para(item)

doc.add_page_break()

# ==============================
# 1. PROJECT OVERVIEW
# ==============================
add_heading('1. Project Overview & Architecture', level=1)

add_para('SolarVista is a full-stack web application that predicts solar panel power output for any city worldwide. '
         'It uses machine learning (XGBoost + RandomForest ensemble) trained on NASA solar radiation data, '
         'combined with real-time weather forecasts from OpenWeatherMap API.')

add_para('')
add_para('Key Features:', bold=True)
add_bullet('Hourly solar radiation predictions for 1, 3, or 5 days')
add_bullet('Panel power output based on user-specified area, efficiency, and tilt angle')
add_bullet('CO₂ savings and environmental impact equivalents (trees, cars, phone charges)')
add_bullet('Interactive Chart.js visualizations (radiation, power, cumulative energy)')
add_bullet('Dark/light theme with modern glassmorphism UI design')
add_bullet('CSV download of forecast data')
add_bullet('City autocomplete and worldwide support via geocoding')

add_para('')
add_para('How It Works (High-Level Flow):', bold=True)
add_numbered('User enters a city name and panel specifications on the web page')
add_numbered('Flask backend geocodes the city to get latitude/longitude via geopy')
add_numbered('Backend fetches current weather + 5-day forecast from OpenWeatherMap API')
add_numbered('Predictor module runs ML model with weather features + solar hour weighting')
add_numbered('Metrics module calculates panel power, energy, CO₂ savings')
add_numbered('Results are returned as JSON to the frontend')
add_numbered('JavaScript displays metric cards, renders Chart.js charts, and populates tables')

doc.add_page_break()

# ==============================
# 2. FILE STRUCTURE
# ==============================
add_heading('2. File Structure', level=1)

add_code("""Solar Prediction Project/
├── main.py                    # Pipeline runner (data → train → predict)
├── requirements.txt           # Python dependencies
├── README.md                  # Project readme
├── data/
│   ├── nasa_solar_data.csv    # Raw NASA POWER solar data
│   ├── clean_solar_dataset.csv # Cleaned dataset for training
│   ├── weather_forecast.csv   # Saved weather forecasts
│   └── solar_predictions.csv  # Saved predictions
├── models/
│   ├── solar_prediction_model.pkl  # Trained ML model (~80MB)
│   └── model_info.pkl              # Feature columns & metrics
├── src/                       # Data pipeline scripts
│   ├── data_collections.py    # NASA data download + weather API
│   ├── prepare_dataset.py     # Data cleaning & feature engineering
│   ├── train_model.py         # Model training (ensemble)
│   └── predict_solar.py       # Standalone prediction script
├── web/                       # Flask web application
│   ├── app.py                 # Flask routes & API endpoints
│   ├── weather_api.py         # OpenWeatherMap API module
│   ├── predictor.py           # ML prediction module
│   ├── metrics.py             # Panel & environmental calculations
│   ├── templates/
│   │   └── index.html         # Main HTML page
│   └── static/
│       ├── style.css          # CSS styles & themes
│       └── app.js             # Frontend JavaScript
├── tests/
│   └── test_metrics.py        # Unit tests for calculations
└── notebook/
    └── analysis.ipynb         # Data analysis notebook""")

doc.add_page_break()

# ==============================
# 3. DATA FLOW DIAGRAM
# ==============================
add_heading('3. Data Flow Diagram', level=1)

add_para('Training Pipeline (One-time):', bold=True)
add_code("""NASA POWER API → data_collections.py → nasa_solar_data.csv
                                    ↓
                         prepare_dataset.py → clean_solar_dataset.csv
                                    ↓
                            train_model.py → solar_prediction_model.pkl""")

add_para('')
add_para('Web Application Flow (Real-time):', bold=True)
add_code("""User Input (city, panel specs)
    ↓
app.py (Flask route: /predict)
    ↓
weather_api.py → OpenWeatherMap API (current + forecast)
    ↓
predictor.py → ML model prediction (hourly W/m²)
    ↓
metrics.py → Panel power, energy, CO₂ calculations
    ↓
JSON Response → app.js → Chart.js charts + metric cards""")

add_para('')
add_concept('Data Flow Between Modules',
    'Each module has a single responsibility. app.py orchestrates the flow: '
    'it calls weather_api.py first to get weather data, then passes that data to predictor.py '
    'for ML predictions, then sends predictions to metrics.py for derived calculations. '
    'The final result is a JSON dictionary sent to the browser.')

doc.add_page_break()

# ==============================
# 4. data_collections.py
# ==============================
add_heading('4. data_collections.py — Data Gathering', level=1)

add_para('Purpose: Downloads 10 years of daily solar radiation data from NASA POWER API '
         'and 5-day weather forecasts from OpenWeatherMap. This is run once to build the training dataset.', bold=True)

add_para('')
add_heading('Imports & Setup', level=2)
add_code("""import requests
import pandas as pd
import os
from geopy.geocoders import Nominatim""")

add_bullet('requests: HTTP library for making API calls to NASA and OpenWeatherMap')
add_bullet('pandas: Data manipulation library; creates DataFrames (table-like structures)')
add_bullet('os: Operating system utilities; here used for creating directories')
add_bullet('geopy.geocoders.Nominatim: Free geocoding service that converts city names to lat/lon coordinates')

add_concept('Geocoding',
    'Geocoding is the process of converting a human-readable address (like "Hyderabad") '
    'into geographic coordinates (latitude: 17.38, longitude: 78.49). The Nominatim service '
    'from OpenStreetMap does this for free.')

add_para('')
add_heading('City Coordinates', level=2)
add_code("""city = "Hyderabad"
geolocator = Nominatim(user_agent="solar_prediction_app")
location = geolocator.geocode(city)
lat = location.latitude
lon = location.longitude""")

add_bullet('city = "Hyderabad": Define the city for which we want solar data')
add_bullet('Nominatim(user_agent=...): Create a geocoder instance. The user_agent is a required identifier.')
add_bullet('geolocator.geocode(city): Sends the city name to Nominatim API, returns a Location object')
add_bullet('location.latitude / .longitude: Extract numeric coordinates from the Location object')

add_para('')
add_heading('NASA POWER API Call', level=2)
add_code("""nasa_url = f"https://power.larc.nasa.gov/api/temporal/daily/point?parameters=ALLSKY_SFC_SW_DWN,T2M,WS2M,RH2M,CLOUD_AMT,PRECTOTCORR,PS&community=RE&longitude={lon}&latitude={lat}&start=20140101&end=20240101&format=JSON"
response = requests.get(nasa_url)
data = response.json()""")

add_bullet('NASA POWER API provides free satellite-derived solar and weather data')
add_bullet('ALLSKY_SFC_SW_DWN: All-sky surface shortwave downward irradiance (solar radiation in kWh/m²/day)')
add_bullet('T2M: Temperature at 2 meters, WS2M: Wind speed at 2m, RH2M: Relative humidity')
add_bullet('CLOUD_AMT: Cloud amount (%), PRECTOTCORR: Precipitation, PS: Surface pressure')
add_bullet('start=20140101&end=20240101: 10 years of daily data (2014-2024)')
add_bullet('requests.get(): Makes HTTP GET request to the API')
add_bullet('.json(): Parses the response from JSON text into a Python dictionary')

add_concept('NASA POWER API',
    'NASA POWER (Prediction of Worldwide Energy Resources) provides meteorological and solar '
    'energy data collected from satellites. It is free and does not require an API key, making it '
    'ideal for training ML models.')

add_para('')
add_heading('Creating DataFrame', level=2)
add_code("""df = pd.DataFrame({
    "date": solar_data.keys(),
    "solar_radiation": solar_data.values(),
    "temperature": temp_data.values(),
    ...
    "latitude": lat,
    "longitude": lon
})
df.to_csv("data/nasa_solar_data.csv", index=False)""")

add_bullet('pd.DataFrame({...}): Creates a table with named columns from dictionary of lists')
add_bullet('solar_data.keys(): Dates are stored as dictionary keys (e.g., "20140101")')
add_bullet('.values(): The actual numeric measurements for each date')
add_bullet('"latitude": lat: A scalar value gets broadcast to every row (same latitude for all data)')
add_bullet('.to_csv(): Saves the DataFrame to a CSV file; index=False omits the row numbers')

add_para('')
add_heading('Weather Forecast Collection', level=2)
add_code("""API_KEY = "11a738ef47063222b2d5e25d33034760"
weather_url = f"https://api.openweathermap.org/data/2.5/forecast?lat={lat}&lon={lon}&appid={API_KEY}&units=metric"
weather_response = requests.get(weather_url)
weather_data = weather_response.json()""")

add_bullet('API_KEY: Your personal OpenWeatherMap API key (free tier)')
add_bullet('"/data/2.5/forecast": 5-day/3-hour forecast endpoint (free tier)')
add_bullet('&units=metric: Returns temperature in °C, wind in m/s')
add_bullet('The response contains a "list" array with weather data at 3-hour intervals')

add_concept('OpenWeatherMap Free Tier',
    'The free tier provides: current weather, 5-day forecast at 3-hour intervals, '
    'geocoding, and 60 API calls per minute. True hourly data requires the paid '
    '"One Call API" subscription.')

doc.add_page_break()

# ==============================
# 5. prepare_dataset.py
# ==============================
add_heading('5. prepare_dataset.py — Data Preprocessing', level=1)

add_para('Purpose: Cleans the raw NASA data, engineers time-based features, and creates '
         'the final dataset used for ML model training.', bold=True)

add_para('')
add_heading('Loading & Date Conversion', level=2)
add_code("""df = pd.read_csv("data/nasa_solar_data.csv")
df["date"] = pd.to_datetime(df["date"], format="%Y%m%d")""")

add_bullet('pd.read_csv(): Reads the CSV file back into a DataFrame')
add_bullet('pd.to_datetime(): Converts string dates like "20140101" to proper DateTime objects')
add_bullet('format="%Y%m%d": Tells pandas the date format (Year-Month-Day with no separators)')

add_para('')
add_heading('Feature Engineering', level=2)
add_code("""df["day_of_year"] = df["date"].dt.dayofyear    # 1-365
df["month"] = df["date"].dt.month              # 1-12
df["hour_of_day"] = 12                         # Daily data = noon average""")

add_bullet('day_of_year: The day number within the year (Jan 1 = 1, Dec 31 = 365). Captures seasonal patterns.')
add_bullet('month: Month number (1-12). Helps the model learn monthly solar patterns.')
add_bullet('hour_of_day = 12: Since NASA data is daily averages, we assign noon (12) as the representative hour. '
           'When predicting for specific hours, the model uses the actual hour to adjust.')

add_concept('Feature Engineering',
    'Feature engineering is the process of creating new input variables from raw data that '
    'help the ML model learn better. day_of_year captures seasonal variation (more sun in summer), '
    'and month captures monthly patterns (monsoon = less sun).')

add_para('')
add_heading('Data Cleaning', level=2)
add_code("""df.replace(-999, None, inplace=True)  # NASA uses -999 for missing
df = df.dropna()                      # Remove rows with any missing values""")

add_bullet('replace(-999, None): NASA POWER uses -999 as a sentinel value for missing/invalid data')
add_bullet('dropna(): Removes any row that has at least one missing (None/NaN) value')
add_bullet('inplace=True: Modifies the DataFrame directly instead of creating a copy')

add_para('')
add_heading('Longitude Support', level=2)
add_code("""if "longitude" not in df.columns:
    df["longitude"] = 78.4867  # Hyderabad default""")

add_bullet('Backward compatibility: if the old NASA data CSV doesn\'t have longitude, add it')
add_bullet('78.4867 is Hyderabad\'s longitude — the city used for training data')

add_concept('Why Longitude Matters',
    'Longitude affects the local solar time. Cities at different longitudes experience solar noon '
    'at different UTC times. The model uses longitude along with hour_of_day to understand '
    'time-based radiation patterns.')

doc.add_page_break()

# ==============================
# 6. train_model.py
# ==============================
add_heading('6. train_model.py — ML Model Training', level=1)

add_para('Purpose: Trains an ensemble ML model (XGBoost + RandomForest) on the cleaned solar dataset '
         'to predict solar radiation from weather and time features.', bold=True)

add_para('')
add_heading('Feature Columns', level=2)
add_code("""FEATURE_COLUMNS = [
    "temperature",       # Air temperature in °C
    "wind_speed",        # Wind speed in m/s
    "humidity",          # Relative humidity in %
    "cloud_cover",       # Cloud coverage in %
    "precipitation",     # Rainfall in mm/day
    "pressure",          # Atmospheric pressure in kPa
    "latitude",          # Geographic latitude
    "longitude",         # Geographic longitude
    "day_of_year",       # 1-365 (seasonal position)
    "month",             # 1-12 (monthly pattern)
    "hour_of_day"        # 0-23 (time of day)
]""")

add_bullet('These 11 features are inputs (X) to the model')
add_bullet('The target variable (y) is solar_radiation (kWh/m²/day)')
add_bullet('Each feature contributes differently: cloud_cover and humidity have strong negative correlations '
           'with solar radiation; day_of_year captures seasonal patterns')

add_para('')
add_heading('Train-Test Split', level=2)
add_code("""X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.2, random_state=42
)""")

add_bullet('test_size=0.2: Reserves 20% of data for testing (unseen data to evaluate model)')
add_bullet('random_state=42: Ensures reproducible splits — same split every time')
add_bullet('X_train/y_train: 80% of data used to teach the model')
add_bullet('X_test/y_test: 20% held out to evaluate how well the model generalizes')

add_concept('Train-Test Split',
    'We split data to prevent overfitting. The model learns patterns only from training data, '
    'then we check its accuracy on the test data (which it has never seen). If the model does well '
    'on test data, it should generalize to real-world predictions.')

add_para('')
add_heading('Ensemble Model', level=2)
add_code("""rf_model = RandomForestRegressor(
    n_estimators=200, max_depth=20, min_samples_split=5,
    random_state=42, n_jobs=-1
)

boost_model = XGBRegressor(
    n_estimators=200, max_depth=8, learning_rate=0.1,
    subsample=0.8, colsample_bytree=0.8, random_state=42
)

model = VotingRegressor(
    estimators=[("rf", rf_model), ("boost", boost_model)]
)""")

add_bullet('RandomForestRegressor: An ensemble of 200 decision trees that vote on the prediction')
add_bullet('  n_estimators=200: Number of trees. More trees = better accuracy but slower')
add_bullet('  max_depth=20: Maximum tree depth (how many decision levels)')
add_bullet('  n_jobs=-1: Use all CPU cores for parallel training')
add_bullet('XGBRegressor: Gradient boosting algorithm that builds trees sequentially, each correcting errors of the previous')
add_bullet('  learning_rate=0.1: How much each new tree contributes (smaller = more conservative)')
add_bullet('  subsample=0.8: Each tree sees 80% of data (reduces overfitting)')
add_bullet('  colsample_bytree=0.8: Each tree sees 80% of features')
add_bullet('VotingRegressor: Combines both models by averaging their predictions')

add_concept('Ensemble Learning',
    'An ensemble combines multiple models to get better predictions than any single model. '
    'RandomForest and XGBoost have different strengths: RandomForest is robust to noise, '
    'while XGBoost captures complex patterns. Averaging them (VotingRegressor) gives the best of both worlds. '
    'Our model achieves R²=0.814 (explains 81.4% of variance in solar radiation).')

add_para('')
add_heading('Model Evaluation', level=2)
add_code("""mae = mean_absolute_error(y_test, predictions)   # Average prediction error
rmse = np.sqrt(mean_squared_error(y_test, predictions))  # Penalizes large errors
r2 = r2_score(y_test, predictions)                       # % of variance explained""")

add_bullet('MAE (Mean Absolute Error) ≈ 0.386: On average, predictions differ by 0.386 kWh/m²/day from actual')
add_bullet('RMSE (Root Mean Squared Error): Similar to MAE but penalizes larger errors more heavily')
add_bullet('R² Score = 0.814: The model explains 81.4% of the variation in solar radiation')

add_para('')
add_heading('Saving the Model', level=2)
add_code("""joblib.dump(model, "models/solar_prediction_model.pkl")
joblib.dump(feature_info, "models/model_info.pkl")""")

add_bullet('joblib.dump(): Serializes the trained model to a binary .pkl file (~80MB)')
add_bullet('The model can be loaded later without retraining using joblib.load()')
add_bullet('model_info.pkl stores metadata: feature column names, accuracy metrics, etc.')

doc.add_page_break()

# ==============================
# 7. predict_solar.py
# ==============================
add_heading('7. predict_solar.py — Standalone Predictions', level=1)

add_para('Purpose: A command-line script that loads the trained model and generates predictions '
         'from weather forecast data. Useful for batch processing without the web interface.', bold=True)

add_para('')
add_code("""model = joblib.load("models/solar_prediction_model.pkl")
forecast_df = pd.read_csv("data/weather_forecast.csv")""")

add_bullet('Loads the pre-trained model from disk')
add_bullet('Reads the weather forecast CSV (previously saved by data_collections.py)')

add_code("""forecast_df["day_of_year"] = forecast_df["datetime"].dt.dayofyear
forecast_df["month"] = forecast_df["datetime"].dt.month
forecast_df["hour_of_day"] = forecast_df["datetime"].dt.hour""")

add_bullet('Engineers the same time features the model was trained on')
add_bullet('hour_of_day: Extracted from forecast timestamps (0-23)')
add_bullet('These features must exactly match what the model expects')

add_concept('Feature Consistency',
    'The prediction input must have the exact same features in the exact same order as the training data. '
    'If the model was trained with 11 features, predictions must provide all 11, named identically.')

doc.add_page_break()

# ==============================
# 8. weather_api.py
# ==============================
add_heading('8. weather_api.py — OpenWeatherMap Integration', level=1)

add_para('Purpose: Centralizes all weather API calls, geocoding, caching, and hourly interpolation. '
         'Used by app.py during real-time predictions.', bold=True)

add_para('')
add_heading('Caching System', level=2)
add_code("""_weather_cache = {}
_CACHE_TTL = 1800  # 30 minutes

def _get_cached(cache_key):
    if cache_key in _weather_cache:
        data, timestamp = _weather_cache[cache_key]
        if time.time() - timestamp < _CACHE_TTL:
            return data
    return None""")

add_bullet('In-memory dictionary cache to avoid redundant API calls')
add_bullet('TTL (Time To Live) = 1800 seconds = 30 minutes')
add_bullet('Each cache entry stores (data, timestamp) tuple')
add_bullet('On lookup: if the entry is less than 30 min old, return cached data; else, delete and re-fetch')

add_concept('API Response Caching',
    'Caching stores API responses temporarily so repeated requests for the same city don\'t '
    'hit the API again. This reduces latency (instant response for cached cities), avoids API '
    'rate limits (60 calls/min on free tier), and improves reliability (cached data available if API is down).')

add_para('')
add_heading('Hourly Interpolation', level=2)
add_code("""def _interpolate_to_hourly(raw_points):
    for i in range(len(raw_points) - 1):
        p1 = raw_points[i]
        p2 = raw_points[i + 1]
        gap_hours = (dt2 - dt1) // 3600
        for h in range(gap_hours):
            frac = h / gap_hours
            interpolated[field] = p1[field] + frac * (p2[field] - p1[field])""")

add_bullet('The free API gives data every 3 hours. We need hourly data for charts.')
add_bullet('Linear interpolation: for each hour between two 3-hour points, calculate intermediate values')
add_bullet('frac = h / gap_hours: Fraction between 0 (at p1) and 1 (at p2)')
add_bullet('interpolated = p1 + frac × (p2 - p1): Linear blend between two data points')
add_bullet('Example: if temperature is 25°C at 12:00 and 28°C at 15:00, then at 13:00 it\'s ≈ 26°C')

add_concept('Linear Interpolation',
    'Linear interpolation estimates a value between two known data points by drawing a straight line '
    'between them. Formula: y = y1 + (x - x1) × (y2 - y1) / (x2 - x1). It\'s simple but effective '
    'for smoothly varying weather parameters.')

add_para('')
add_heading('City Autocomplete', level=2)
add_code("""POPULAR_CITIES = [
    "Mumbai", "Delhi", "Bangalore", "Hyderabad", ...
    "New York", "London", "Tokyo", "Sydney", "Dubai", ...
]

def search_cities(query):
    query_lower = query.lower()
    return [c for c in POPULAR_CITIES if query_lower in c.lower()][:10]""")

add_bullet('Pre-defined list of 45+ popular cities worldwide for quick autocomplete')
add_bullet('Case-insensitive substring matching: typing "mum" matches "Mumbai"')
add_bullet('Returns max 10 results to keep the dropdown manageable')
add_bullet('Note: Users can type ANY city — autocomplete is just for convenience. Any city typed '
           'will be geocoded by Nominatim, which supports all cities worldwide.')

doc.add_page_break()

# ==============================
# 9. predictor.py
# ==============================
add_heading('9. predictor.py — ML Prediction Module', level=1)

add_para('Purpose: Handles ML model loading, solar radiation prediction, and hourly distribution '
         'using a solar bell-curve to convert daily averages to realistic hourly W/m² values.', bold=True)

add_para('')
add_heading('Daylight Hours Calculation', level=2)
add_code("""def _get_daylight_hours(lat, day_of_year):
    declination = 23.45 * math.sin(math.radians(360/365 * (284 + day_of_year)))
    cos_hour = -math.tan(lat_rad) * math.tan(decl_rad)
    hour_angle = math.degrees(math.acos(cos_hour))
    daylight = 2 * hour_angle / 15""")

add_bullet('declination: The tilt of Earth\'s axis relative to the sun (varies -23.45° to +23.45° over the year)')
add_bullet('Determines sunrise/sunset times based on latitude and season')
add_bullet('At equator (lat=0): ~12 hours daylight year-round')
add_bullet('At higher latitudes: more variation (long summer days, short winter days)')
add_bullet('This is used to determine when solar radiation is non-zero')

add_concept('Solar Declination',
    'Earth\'s axis is tilted 23.45° from the orbital plane. As Earth orbits the sun, '
    'different latitudes receive different amounts of sunlight. The declination angle '
    'tells us the sun\'s position relative to the equator on any given day. '
    'June 21 = +23.45° (longest day in North), December 21 = -23.45° (shortest day).')

add_para('')
add_heading('Solar Hour Weighting (Bell Curve)', level=2)
add_code("""def _solar_hour_weight(hour, sunrise_hour, sunset_hour):
    if hour < sunrise_hour or hour > sunset_hour:
        return 0
    solar_noon = (sunrise_hour + sunset_hour) / 2
    normalized = (hour - solar_noon) / half_day
    weight = math.cos(normalized * math.pi / 2) ** 2""")

add_bullet('Returns 0 for nighttime hours (before sunrise or after sunset)')
add_bullet('Returns maximum weight at solar noon (peak sun)')
add_bullet('Uses cos² shape — realistic bell curve matching actual solar radiation patterns')
add_bullet('Example: if sunrise=6:00, sunset=18:00, noon=12:00: weight at 12:00 = 1.0, at 9:00 ≈ 0.75, at 6:00 = 0')

add_concept('Solar Bell Curve',
    'Solar radiation doesn\'t jump on/off at sunrise/sunset — it gradually increases from sunrise, '
    'peaks at solar noon, and gradually decreases until sunset. A cosine-squared function closely '
    'models this real-world pattern. The model predicts DAILY average radiation; the bell curve '
    'distributes it across hours realistically.')

add_para('')
add_heading('Converting Daily to Hourly W/m²', level=2)
add_code("""daily_avg = model.predict(data)[0]  # kWh/m²/day
radiation_wm2 = daily_avg * 1000 / daylight_hours * peak_factor * weight""")

add_bullet('model.predict(): Returns daily average solar radiation in kWh/m²/day (NASA\'s unit)')
add_bullet('× 1000: Convert kWh to Wh')
add_bullet('÷ daylight_hours: Distribute over sunlit hours')
add_bullet('× peak_factor (2.0): Accounts for bell curve shape (the integral of cos² is half the period)')
add_bullet('× weight: Apply the hour-specific bell curve weight')
add_bullet('Result: instantaneous W/m² at the specific hour, matching real solar patterns')

doc.add_page_break()

# ==============================
# 10. metrics.py
# ==============================
add_heading('10. metrics.py — Panel & Environmental Calculations', level=1)

add_para('Purpose: Computes all derived metrics from solar radiation predictions — panel output, '
         'energy production, CO₂ savings, and environmental equivalents.', bold=True)

add_para('')
add_heading('Panel Irradiance', level=2)
add_code("""def calculate_panel_irradiance(radiation_wm2, tilt_deg):
    tilt_rad = math.radians(tilt_deg)
    return radiation_wm2 * math.cos(tilt_rad)""")

add_bullet('Converts tilt angle from degrees to radians (math functions use radians)')
add_bullet('Applies cosine correction: tilted panels receive less direct sunlight')
add_bullet('At 0° tilt (flat): cos(0) = 1.0, full radiation received')
add_bullet('At 30° tilt: cos(30°) = 0.866, panel receives 86.6% of direct radiation')
add_bullet('At 90° tilt (vertical): cos(90°) ≈ 0, almost no direct radiation')

add_para('')
add_heading('Panel Power Output', level=2)
add_code("""def calculate_panel_power(radiation_wm2, area_m2, efficiency_pct, tilt_deg):
    irradiance = calculate_panel_irradiance(radiation_wm2, tilt_deg)
    efficiency = efficiency_pct / 100.0
    power = irradiance * area_m2 * efficiency
    return max(0, round(power, 2))""")

add_bullet('Formula: Power (W) = Irradiance (W/m²) × Area (m²) × Efficiency')
add_bullet('efficiency_pct / 100: Converts percentage to decimal (18% → 0.18)')
add_bullet('max(0, ...): Ensures power is never negative')
add_bullet('Example: 800 W/m² × 10 m² × 0.18 efficiency × cos(30°) = 800 × 10 × 0.18 × 0.866 = 1246.1 W')

add_concept('Solar Panel Efficiency',
    'Solar panel efficiency is the percentage of incoming solar energy that the panel converts to electricity. '
    'Typical values: monocrystalline = 18-22%, polycrystalline = 15-17%, thin-film = 10-13%. '
    'Higher efficiency means more power from the same panel area, but costs more.')

add_para('')
add_heading('CO₂ Savings', level=2)
add_code("""def calculate_co2_savings(energy_kwh):
    return round(energy_kwh * 0.82, 3)""")

add_bullet('Grid emission factor: 0.82 kg CO₂ per kWh (India average)')
add_bullet('Every kWh of solar energy replaces 0.82 kg of CO₂ that would have been emitted by coal/gas power plants')
add_bullet('Example: 5 kWh solar → 4.1 kg CO₂ avoided')

add_para('')
add_heading('Environmental Equivalents', level=2)
add_code("""def calculate_environmental_equivalents(co2_savings_kg):
    return {
        "trees_planted": co2_savings_kg / 21.77,     # kg CO₂/tree/year
        "cars_off_road": co2_savings_kg / 12.0,       # kg CO₂/car/day
        "phone_charges": co2_savings_kg / 0.008,      # kg CO₂/charge
        "led_bulb_hours": co2_savings_kg / 0.0082,    # kg CO₂/hour
    }""")

add_bullet('trees_planted: 1 mature tree absorbs ~21.77 kg CO₂ per year (EPA estimate)')
add_bullet('cars_off_road: Average car emits ~12 kg CO₂ per day (US average)')
add_bullet('phone_charges: One smartphone charge ≈ 8g CO₂ (including grid losses)')
add_bullet('led_bulb_hours: A 10W LED bulb running for 1 hour ≈ 8.2g CO₂ from grid')

doc.add_page_break()

# ==============================
# 11. app.py
# ==============================
add_heading('11. app.py — Flask Backend & Routes', level=1)

add_para('Purpose: The Flask web server that handles HTTP requests, orchestrates the prediction '
         'pipeline, and serves the web interface.', bold=True)

add_para('')
add_heading('Flask Application Setup', level=2)
add_code("""from flask import Flask, render_template, request, jsonify, Response
app = Flask(__name__)""")

add_bullet('Flask: A lightweight Python web framework')
add_bullet('render_template: Renders HTML templates with Jinja2 templating')
add_bullet('request: Access incoming HTTP request data (form data, JSON, query params)')
add_bullet('jsonify: Converts Python dicts to JSON HTTP responses')
add_bullet('Response: Create custom HTTP responses (used for CSV download)')

add_concept('Flask',
    'Flask is a micro web framework for Python. It maps URL routes to Python functions. '
    'When a browser requests a URL, Flask calls the corresponding function and returns its output '
    'as an HTTP response. "Micro" means it provides the essentials without imposing decisions on tools.')

add_para('')
add_heading('Route: Home Page (GET /)', level=2)
add_code("""@app.route("/")
def home():
    return render_template("index.html")""")

add_bullet('@app.route("/"): Decorator that maps the root URL to this function')
add_bullet('render_template("index.html"): Loads templates/index.html and returns it as HTML')
add_bullet('This serves the initial page load — all dynamic content is loaded via JavaScript afterwards')

add_para('')
add_heading('Route: Predict (POST /predict)', level=2)
add_code("""@app.route("/predict", methods=["POST"])
def predict():
    data = request.get_json()
    city = data.get("city", "").strip()""")

add_bullet('methods=["POST"]: Only accepts POST requests (form submissions with data)')
add_bullet('request.get_json(): Parses the JSON body sent by the frontend JavaScript')
add_bullet('.get("city", ""): Safely retrieves the "city" key, defaulting to empty string if missing')
add_bullet('.strip(): Removes leading/trailing whitespace from the city name')

add_para('')
add_heading('Input Validation', level=2)
add_code("""try:
    panel_area = float(data.get("panel_area", 10))
    panel_efficiency = float(data.get("panel_efficiency", 18))
    panel_tilt = float(data.get("panel_tilt", 30))
except (ValueError, TypeError):
    return jsonify({"error": "Invalid numeric input"}), 400

if not (0.1 <= panel_area <= 10000):
    return jsonify({"error": "Panel area must be between 0.1 and 10000 m²"}), 400""")

add_bullet('try/except: Catches invalid number conversions (e.g., user sends "abc" as panel area)')
add_bullet('Range validation: Ensures values are physically realistic')
add_bullet('return jsonify({...}), 400: Returns a JSON error with HTTP 400 (Bad Request) status')

add_concept('HTTP Status Codes',
    '200 = Success (OK), 400 = Bad Request (invalid input), 404 = Not Found (city doesn\'t exist), '
    '500 = Server Error (unexpected bug), 502 = Bad Gateway (external API failed). '
    'The frontend checks these codes to show appropriate error messages.')

add_para('')
add_heading('Prediction Pipeline', level=2)
add_code("""location = geocode_city(city)           # Get lat/lon
current_weather = get_current_weather(lat, lon)  # Current conditions
forecast = get_forecast(lat, lon)          # 5-day forecast
hourly_predictions = predict_hourly(forecast, lat, lon)  # ML predictions
full_metrics = compute_full_metrics(hourly_predictions, ...)  # Derived metrics
return jsonify(full_metrics)               # JSON response to browser""")

add_bullet('This is the core pipeline: geocode → weather → predict → metrics → respond')
add_bullet('Each function call delegates to the appropriate module')
add_bullet('If any step fails (invalid city, API down), it returns an error JSON before reaching subsequent steps')
add_bullet('jsonify(): Converts the Python dict to a JSON string with proper Content-Type header')

add_para('')
add_heading('Route: CSV Download (POST /download/csv)', level=2)
add_code("""@app.route("/download/csv", methods=["POST"])
def download_csv():
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["DateTime", "Radiation (W/m²)", ...])
    return Response(csv_content, mimetype="text/csv",
                    headers={"Content-Disposition": "attachment; filename=..."})""")

add_bullet('io.StringIO(): Creates an in-memory text stream (acts like a file but lives in RAM)')
add_bullet('csv.writer(): Writes CSV-formatted data to the stream')
add_bullet('Content-Disposition: attachment: Tells the browser to download the file instead of displaying it')
add_bullet('The mimetype="text/csv" tells the browser this is a CSV file')

doc.add_page_break()

# ==============================
# 12. index.html
# ==============================
add_heading('12. index.html — Frontend Layout', level=1)

add_para('Purpose: The single-page HTML that structures the entire user interface — input form, '
         'result cards, charts, and download functionality.', bold=True)

add_para('')
add_heading('Document Head', level=2)
add_code("""<html lang="en" data-theme="dark">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>SolarVista — Solar Power Prediction</title>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800" rel="stylesheet">
    <link rel="stylesheet" href="{{ url_for('static', filename='style.css') }}">
    <script src="https://cdn.jsdelivr.net/npm/chart.js@4.4.0/dist/chart.umd.min.js"></script>
</head>""")

add_bullet('data-theme="dark": Custom HTML attribute used by CSS to switch themes')
add_bullet('viewport meta: Enables responsive design on mobile devices')
add_bullet('Google Fonts Inter: Modern, clean font loaded from Google CDN')
add_bullet('{{ url_for(...) }}: Jinja2 template function that generates the correct URL for static files')
add_bullet('Chart.js: JavaScript charting library loaded from CDN')

add_concept('Jinja2 Templating',
    'Flask uses Jinja2 to inject dynamic content into HTML. The {{ ... }} syntax outputs a value, '
    'and {% ... %} is for control flow (if/for). url_for() generates URLs relative to the app '
    'root, ensuring paths work regardless of deployment configuration.')

add_para('')
add_heading('Input Form with Sliders', level=2)
add_code("""<input type="range" id="panel-area" name="panel_area"
       min="1" max="500" value="10" step="1">""")

add_bullet('type="range": Creates a horizontal slider instead of text input')
add_bullet('min/max/step: Constrain the slider range and increment')
add_bullet('value="10": Default starting value')
add_bullet('JavaScript updates the displayed value label as the slider moves')

add_para('')
add_heading('Metric Cards (Results)', level=2)
add_code("""<div class="metric-card glass-card solar-card" id="card-radiation">
    <div class="metric-icon">☀️</div>
    <div class="metric-value" id="val-radiation">—</div>
    <div class="metric-label">Solar Radiation</div>
    <div class="metric-unit">W/m²</div>
</div>""")

add_bullet('glass-card: CSS class that applies glassmorphism effect (blurred background)')
add_bullet('metric-value: The actual number, updated by JavaScript after prediction')
add_bullet('Initially shows "—" (em dash) as placeholder until data arrives')
add_bullet('4 cards: Solar Radiation (W/m²), Panel Power (W), Total Energy (kWh), CO₂ Saved (kg)')

add_para('')
add_heading('Chart.js Canvases', level=2)
add_code("""<canvas id="radiation-chart"></canvas>
<canvas id="power-chart"></canvas>
<canvas id="energy-chart"></canvas>""")

add_bullet('HTML <canvas> elements are where Chart.js draws the charts')
add_bullet('JavaScript later creates Chart objects targeting these canvas elements by ID')
add_bullet('3 charts: hourly radiation, hourly power output, cumulative energy')

add_concept('HTML Canvas',
    'The <canvas> element provides a blank drawing surface for JavaScript. '
    'Libraries like Chart.js use the Canvas API to render pixels directly, enabling smooth '
    'animations and interactive charts that HTML tables cannot provide.')

doc.add_page_break()

# ==============================
# 13. style.css
# ==============================
add_heading('13. style.css — Styling & Themes', level=1)

add_para('Purpose: Complete CSS design system with dark/light themes, glassmorphism effects, '
         'responsive layouts, and micro-animations.', bold=True)

add_para('')
add_heading('CSS Custom Properties (Variables)', level=2)
add_code(""":root {
    --bg-primary: #0a0e1a;
    --bg-card: rgba(255, 255, 255, 0.04);
    --text-primary: #f0f0f5;
    --accent-solar: #ffb347;
    --blur-amount: 20px;
}

[data-theme="light"] {
    --bg-primary: #f5f7ff;
    --bg-card: rgba(255, 255, 255, 0.7);
    --text-primary: #1a1a2e;
}""")

add_bullet(':root: Special CSS selector targeting the document root (html element)')
add_bullet('--variable-name: CSS custom properties (variables) that can be referenced with var()')
add_bullet('[data-theme="light"]: CSS attribute selector — when the HTML has data-theme="light", '
           'these variables override the defaults')
add_bullet('Switching themes = changing one HTML attribute; all colors update automatically via CSS variables')

add_concept('CSS Custom Properties',
    'CSS variables (--name) allow defining values once and using them everywhere with var(--name). '
    'When the value changes, all elements using that variable update automatically. This is how '
    'our dark/light theme works: JavaScript changes data-theme, CSS variables swap, every color updates.')

add_para('')
add_heading('Glassmorphism Effect', level=2)
add_code(""".glass-card {
    background: var(--bg-card);              /* Semi-transparent background */
    backdrop-filter: blur(var(--blur-amount)); /* Blur content behind the card */
    border: 1px solid var(--border-card);     /* Subtle border */
    border-radius: 20px;                      /* Rounded corners */
    box-shadow: var(--shadow-card);           /* Soft shadow for depth */
}""")

add_bullet('background: rgba(255,255,255,0.04): Nearly transparent white — lets background show through')
add_bullet('backdrop-filter: blur(20px): Blurs the content BEHIND the card (frosted glass effect)')
add_bullet('-webkit-backdrop-filter: Same property for Safari compatibility')
add_bullet('border-radius: 20px: Rounded corners for modern look')
add_bullet('Combined: creates a frosted glass surface floating over the gradient background')

add_concept('Glassmorphism',
    'A modern UI design trend featuring semi-transparent elements with blur effects, creating '
    'the appearance of frosted glass. It provides depth and visual hierarchy while maintaining '
    'the aesthetic of the background gradient. Popularized by Apple\'s iOS and macOS design.')

add_para('')
add_heading('Animated Background Orbs', level=2)
add_code("""body::before {
    content: '';
    width: 500px; height: 500px;
    background: var(--accent-solar);
    border-radius: 50%;
    filter: blur(80px);
    opacity: 0.15;
    animation: float-orb 20s ease-in-out infinite;
}""")

add_bullet('::before pseudo-element: Creates an invisible element as the first child of body')
add_bullet('Large blurred circle (500px) positioned behind all content')
add_bullet('filter: blur(80px): Extremely blurred, creating a soft ambient glow')
add_bullet('opacity: 0.15: Very subtle — adds atmosphere without distracting')
add_bullet('animation: Slowly moves the orb for a living, dynamic feel')

add_para('')
add_heading('Responsive Design', level=2)
add_code("""@media (max-width: 768px) {
    .form-row {
        grid-template-columns: 1fr;  /* Stack vertically on mobile */
    }
    .metrics-grid {
        grid-template-columns: repeat(2, 1fr);  /* 2 columns instead of 4 */
    }
}""")

add_bullet('@media: CSS media query — applies rules only when screen width ≤ 768px (tablets)')
add_bullet('grid-template-columns: 1fr: Single column layout (stacked)')
add_bullet('Three breakpoints: 1024px (tablet), 768px (small tablet), 480px (phone)')

add_concept('Responsive Design',
    'Making websites work well on all screen sizes. CSS Grid and media queries let us '
    'define different layouts for different screen widths. On desktop: 4-column card grid. '
    'On tablet: 2 columns. On phone: 1 column. The content adapts without separate mobile pages.')

doc.add_page_break()

# ==============================
# 14. app.js
# ==============================
add_heading('14. app.js — Frontend JavaScript Logic', level=1)

add_para('Purpose: Handles all frontend interactivity — API calls, Chart.js rendering, '
         'theme toggle, city autocomplete, slider updates, and CSV download.', bold=True)

add_para('')
add_heading('Fetch API — Prediction Request', level=2)
add_code("""const res = await fetch('/predict', {
    method: 'POST',
    headers: { 'Content-Type': 'application/json' },
    body: JSON.stringify(payload)
});
const data = await res.json();""")

add_bullet('fetch(): Modern browser API for making HTTP requests (replaces older XMLHttpRequest)')
add_bullet('method: "POST": Sends data to the server (vs GET which only requests data)')
add_bullet('Content-Type: application/json: Tells the server we\'re sending JSON')
add_bullet('JSON.stringify(payload): Converts JavaScript object to JSON string')
add_bullet('await res.json(): Parses the JSON response body into a JavaScript object')
add_bullet('await: Pauses execution until the async operation completes (avoids callback hell)')

add_concept('Fetch API & async/await',
    'fetch() is the modern way to make HTTP requests from JavaScript. Combined with async/await, '
    'it reads like synchronous code while being non-blocking (the browser stays responsive). '
    'The function pauses at "await" until the network request completes, then continues with the result.')

add_para('')
add_heading('Chart.js Rendering', level=2)
add_code("""radiationChart = new Chart(document.getElementById('radiation-chart'), {
    type: 'line',
    data: {
        labels: labels,
        datasets: [{
            label: 'Solar Radiation (W/m²)',
            data: radiationData,
            borderColor: '#ffb347',
            backgroundColor: 'rgba(255,179,71,0.1)',
            fill: true,
            tension: 0.4,
        }],
    },
    options: { ... }
});""")

add_bullet('new Chart(): Creates a Chart.js chart instance on the specified canvas element')
add_bullet('type: "line": Creates a line chart (other types: bar, pie, doughnut, etc.)')
add_bullet('labels: Array of x-axis labels (timestamps like "Mar 15, 2PM")')
add_bullet('data: Array of y-axis values (radiation values in W/m²)')
add_bullet('borderColor: Line color, backgroundColor: Fill color under the line')
add_bullet('fill: true: Fills the area under the line with semi-transparent color')
add_bullet('tension: 0.4: Smooth curved lines (0 = straight, 1 = very curved)')

add_concept('Chart.js',
    'Chart.js is an open-source JavaScript charting library that renders interactive charts '
    'on HTML canvas. It supports line, bar, pie, radar, and many other chart types with '
    'built-in animations, tooltips, and responsive sizing. Loaded via CDN — no installation needed.')

add_para('')
add_heading('Theme Toggle with localStorage', level=2)
add_code("""const savedTheme = localStorage.getItem('solarvista-theme') || 'dark';
document.documentElement.setAttribute('data-theme', savedTheme);

themeToggle.addEventListener('click', () => {
    const next = current === 'dark' ? 'light' : 'dark';
    document.documentElement.setAttribute('data-theme', next);
    localStorage.setItem('solarvista-theme', next);
});""")

add_bullet('localStorage: Browser storage that persists between page loads (like cookies but simpler)')
add_bullet('getItem()/setItem(): Read/write key-value pairs to localStorage')
add_bullet('document.documentElement: The <html> element')
add_bullet('setAttribute("data-theme", ...): Changes the HTML attribute that CSS responds to')
add_bullet('When theme changes, CSS variables update automatically, and charts are re-rendered')

add_para('')
add_heading('Count-Up Animation', level=2)
add_code("""function animateValue(elementId, value, decimals) {
    const duration = 800;
    function update(currentTime) {
        const progress = Math.min(elapsed / duration, 1);
        const eased = 1 - Math.pow(1 - progress, 3);  // ease-out cubic
        const current = start + (value - start) * eased;
        element.textContent = current.toFixed(decimals);
        if (progress < 1) requestAnimationFrame(update);
    }
    requestAnimationFrame(update);
}""")

add_bullet('Animates metric values from 0 to their final value over 800ms')
add_bullet('requestAnimationFrame(): Browser API that calls a function before each screen repaint (60fps)')
add_bullet('Ease-out cubic: Animation starts fast and slows down at the end (feels natural)')
add_bullet('toFixed(decimals): Formats number to specific decimal places (e.g., 3.14)')

add_concept('requestAnimationFrame',
    'requestAnimationFrame is the browser\'s way of scheduling smooth animations. It calls your '
    'function right before the next screen repaint (~60 times per second). This is more efficient '
    'than setInterval because the browser can optimize rendering and pause when the tab is hidden.')

add_para('')
add_heading('City Autocomplete with Debouncing', level=2)
add_code("""let debounceTimer;
cityInput.addEventListener('input', () => {
    clearTimeout(debounceTimer);
    debounceTimer = setTimeout(async () => {
        const res = await fetch(`/cities?q=${query}`);
        const cities = await res.json();
        // Render dropdown suggestions
    }, 200);
});""")

add_bullet('addEventListener("input"): Fires every time the user types a character')
add_bullet('clearTimeout/setTimeout: Debouncing — waits 200ms after the last keystroke before fetching')
add_bullet('Without debouncing: typing "Mumbai" would make 6 API calls (M, Mu, Mum, Mumb, Mumba, Mumbai)')
add_bullet('With debouncing: only 1 call after the user stops typing for 200ms')

add_concept('Debouncing',
    'A technique to limit how often a function runs. Instead of calling it on every event, '
    'we wait until the events "settle down." Each new event resets the timer. Only when no '
    'events occur for the specified delay does the function actually execute.')

doc.add_page_break()

# ==============================
# 15. API LIMITATIONS
# ==============================
add_heading('15. API Limitations & Assumptions', level=1)

add_para('OpenWeatherMap Free Tier:', bold=True)
add_bullet('Provides 5-day forecast at 3-hour intervals (not true hourly)')
add_bullet('We interpolate to hourly using linear interpolation (good approximation for weather)')
add_bullet('Rate limit: 60 API calls per minute')
add_bullet('True hourly data requires "One Call API" subscription ($)')
add_bullet('Geocoding: resolves city names to coordinates (included in free tier)')

add_para('')
add_para('NASA POWER API:', bold=True)
add_bullet('Provides daily average solar radiation (not hourly)')
add_bullet('We use a solar bell curve to distribute daily values across hours')
add_bullet('Data available from 1981 to present, ~1-2 months delay')
add_bullet('Free, no API key required')

add_para('')
add_para('Model Assumptions:', bold=True)
add_bullet('Trained on Hyderabad (17.4°N, 78.5°E) data — may be less accurate for extreme latitudes')
add_bullet('Solar radiation is daily average in kWh/m²/day (NASA unit)')
add_bullet('Bell curve distribution to hourly is an approximation — real patterns vary with elevation and terrain')
add_bullet('Panel efficiency is constant (real panels lose efficiency at high temperatures)')
add_bullet('CO₂ emission factor of 0.82 kg/kWh is India\'s grid average (varies by country)')

doc.add_page_break()

# ==============================
# 16. KEY FORMULAS
# ==============================
add_heading('16. Key Formulas & Calculations', level=1)

add_para('')
add_para('Panel Irradiance:', bold=True)
add_para('  Irradiance (W/m²) = Radiation × cos(tilt_angle)')
add_para('')

add_para('Panel Power Output:', bold=True)
add_para('  Power (W) = Irradiance (W/m²) × Area (m²) × Efficiency (%/100)')
add_para('')

add_para('Hourly Energy:', bold=True)
add_para('  Energy (Wh) = Power (W) × 1 hour')
add_para('')

add_para('Daily Energy:', bold=True)
add_para('  Daily Energy (Wh) = Σ(Hourly Energy values)')
add_para('')

add_para('CO₂ Savings:', bold=True)
add_para('  CO₂ (kg) = Energy (kWh) × 0.82')
add_para('')

add_para('Environmental Equivalents:', bold=True)
add_para('  Trees ≡ CO₂ / 21.77 kg/tree/year')
add_para('  Cars ≡ CO₂ / 12.0 kg/car/day')
add_para('  Phone charges ≡ CO₂ / 0.008 kg/charge')
add_para('  LED hours ≡ CO₂ / 0.0082 kg/hour')

add_para('')
add_para('Daylight Hours:', bold=True)
add_para('  Declination = 23.45 × sin(360/365 × (284 + day_of_year))')
add_para('  Hour angle = arccos(-tan(lat) × tan(declination))')
add_para('  Daylight = 2 × hour_angle / 15')

add_para('')
add_para('Solar Hour Weight (Bell Curve):', bold=True)
add_para('  weight = cos²((hour - solar_noon) / half_day × π/2)')
add_para('  where solar_noon = (sunrise + sunset) / 2')

# ==============================
# SAVE
# ==============================
output_path = os.path.join(
    r"d:\1M1B Internship\Solar Prediction Project",
    "SolarVista_Project_Documentation.docx"
)
doc.save(output_path)
print(f"Document saved to: {output_path}")
